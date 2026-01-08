from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Título principal
title = doc.add_heading('CORREOS ELECTRÓNICOS - TerraTokenX', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Este documento contiene los textos de todos los correos electrónicos que envía el sistema.')
doc.add_paragraph('Los campos entre {{ }} son variables que se reemplazan automáticamente con los datos del cliente.')

# CORREO 1
doc.add_heading('1. CORREO: RESERVA PENDIENTE (MERCADO PAGO)', level=1)
doc.add_paragraph('Cuándo se envía: Cuando el cliente crea una reserva y selecciona Mercado Pago como método de pago, pero aún no ha pagado.')
doc.add_paragraph('Asunto: Tu solicitud de reserva está pendiente de pago').bold = True

doc.add_heading('Contenido:', level=2)
content1 = """
Hola {{ reserva.nombre }},

Hemos recibido tu solicitud de reserva para la Preventa de TerraTokenX.

Tu orden #{{ reserva.numero_reserva }} ha sido creada exitosamente. Para asegurar tu cupo, por favor completa tu pago a través de Mercado Pago.

Resumen de la Orden:
• Monto Total: ${{ reserva.total }}
• Método de Pago: Mercado Pago

Si ya realizaste el pago, recibirás un correo de confirmación en breve.

Saludos,
Equipo TerraTokenX
"""
doc.add_paragraph(content1)

# CORREO 2
doc.add_heading('2. CORREO: RESERVA PENDIENTE (CRYPTO)', level=1)
doc.add_paragraph('Cuándo se envía: Cuando el cliente crea una reserva y selecciona Criptomoneda como método de pago.')
doc.add_paragraph('Asunto: ⏳ Instrucciones para finalizar tu inversión en TerraTokenX').bold = True

doc.add_heading('Contenido:', level=2)
content2 = """
Hola {{ reserva.nombre }},

Hemos recibido tu solicitud de reserva para la Preventa de TerraTokenX.

Para confirmar tu cupo, por favor completa la transferencia del monto acordado a nuestra billetera oficial.

Resumen de Orden:
• Orden: #{{ reserva.numero_reserva }}
• Monto a transferir: ${{ reserva.total }} (Equivalente en Cripto)

Datos de Transferencia (CryptoMarket):
• Wallet (USDT/BTC): [AQUÍ VA LA WALLET QUE TE DARÁ JOAN]
• Red: [AQUÍ VA LA RED, EJ: TRC20]

Importante: Una vez realizada la transferencia, responde a este correo adjuntando el comprobante o Hash de la transacción para validar tu pago manualmente.

Quedamos atentos a tu comprobante.

Saludos,
Equipo TerraTokenX
"""
doc.add_paragraph(content2)

# CORREO 3
doc.add_heading('3. CORREO: CONFIRMACIÓN DE PAGO', level=1)
doc.add_paragraph('Cuándo se envía: Cuando el pago ha sido confirmado (ya sea automáticamente por Mercado Pago o manualmente por el administrador).')
doc.add_paragraph('Asunto: ✅ Confirmación: Tu cupo en la Preventa TerraTokenX está asegurado').bold = True

doc.add_heading('Contenido:', level=2)
content3 = """
Hola {{ reserva.nombre }},

¡Bienvenido a TerraTokenX!

Confirmamos la recepción de tu pago por ${{ reserva.total }}.

Has asegurado exitosamente tu posición en nuestra Preventa Exclusiva.

Detalles de tu Inversión:
• Nº de Orden: #{{ reserva.numero_reserva }}
• Medio de Pago: {{ reserva.metodo_pago }}
• Estado: Confirmado / Pagado

¿Qué sigue ahora? Tu participación ya está registrada en nuestra base de datos prioritaria. En Marzo, cuando lancemos la plataforma oficial, te contactaremos a este mismo correo para migrar tu saldo y entregarte tus Tokens digitales en tu dashboard personal.

Gracias por confiar en el futuro de la inversión inmobiliaria.

Si tienes dudas, contáctanos a nuestro WhatsApp: +56 9 2883 9093

Atentamente,
El equipo de TerraTokenX
Respaldo Legal y Tecnología Blockchain
"""
doc.add_paragraph(content3)

# NOTAS
doc.add_heading('NOTAS PARA EDICIÓN', level=1)
doc.add_paragraph('Variables Dinámicas:')
variables = """
• {{ reserva.nombre }} → Nombre del cliente
• {{ reserva.numero_reserva }} → Número de orden (ej: ORD-2024-001)
• {{ reserva.total }} → Monto total en pesos
• {{ reserva.metodo_pago }} → Mercado Pago o Crypto
"""
doc.add_paragraph(variables)

# Guardar
doc.save('Correos_TerraTokenX.docx')
print("✅ Documento Word creado: Correos_TerraTokenX.docx")
